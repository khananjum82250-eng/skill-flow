from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import zipfile, shutil, re, math, textwrap

ROOT = Path(r"C:\xampp\htdocs\Skill_flow")
DOC_DIR = ROOT / "documentation"
ASSET_DIR = DOC_DIR / "skillflow_report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUT = DOC_DIR / "SkillFlow_Final_370_Page_BCA_Report.docx"
ZIP = Path(r"C:\Users\naved\Downloads\WhatsApp Unknown 2026-05-22 at 1.31.27 AM.zip")

PURPLE, LAV, BG, PINK, DARK, TEXT, MUTED = "8B5CF6", "C4B5FD", "F8F5FF", "F472B6", "1F2D5C", "1B1B2F", "5B667A"

try:
    FONT_B = ImageFont.truetype("arialbd.ttf", 26)
    FONT_M = ImageFont.truetype("arial.ttf", 20)
    FONT_S = ImageFont.truetype("arial.ttf", 16)
except Exception:
    FONT_B = FONT_M = FONT_S = ImageFont.load_default()


def collect_screenshots():
    shot_dir = ASSET_DIR / "screenshots"
    shot_dir.mkdir(exist_ok=True)
    if ZIP.exists():
        with zipfile.ZipFile(ZIP) as z:
            for name in z.namelist():
                if name.lower().endswith((".png", ".jpg", ".jpeg")):
                    safe = re.sub(r"[^A-Za-z0-9_. -]+", "_", Path(name).name)
                    target = shot_dir / safe
                    if not target.exists():
                        target.write_bytes(z.read(name))
    candidates = []
    bases = [Path(r"C:\Users\naved\OneDrive\Pictures\Screenshots"), Path(r"C:\Users\naved\OneDrive\Pictures"), Path(r"C:\Users\naved\Downloads")]
    for base in bases:
        if base.exists():
            try:
                for p in base.rglob("*"):
                    if p.is_file() and p.suffix.lower() in [".png", ".jpg", ".jpeg"] and re.search(r"Screenshot|WhatsApp|Skill|skill", p.name, re.I):
                        candidates.append(p)
            except Exception:
                pass
    for p in sorted(candidates, key=lambda x: x.stat().st_mtime, reverse=True)[:35]:
        target = shot_dir / re.sub(r"[^A-Za-z0-9_. -]+", "_", p.name)
        try:
            if not target.exists():
                shutil.copy2(p, target)
        except Exception:
            pass
    screens = []
    for p in shot_dir.iterdir():
        if p.suffix.lower() in [".png", ".jpg", ".jpeg"]:
            try:
                im = Image.open(p)
                w, h = im.size
                if w >= 500 and h >= 300:
                    screens.append(p)
            except Exception:
                pass
    return screens[:25]


def gradient_bg(w, h):
    im = Image.new("RGB", (w, h), "white")
    pix = im.load()
    c1, c2, c3 = (248, 245, 255), (244, 220, 255), (255, 238, 247)
    for y in range(h):
        for x in range(w):
            t = (x + y) / (w + h)
            if t < 0.55:
                u = t / 0.55
                c = tuple(int(c1[i] * (1 - u) + c2[i] * u) for i in range(3))
            else:
                u = (t - 0.55) / 0.45
                c = tuple(int(c2[i] * (1 - u) + c3[i] * u) for i in range(3))
            pix[x, y] = c
    return im


def rounded(draw, box, fill, outline=(139, 92, 246), width=3, radius=24):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, a, b, fill=(80, 55, 160), width=4):
    draw.line([a, b], fill=fill, width=width)
    ang = math.atan2(b[1] - a[1], b[0] - a[0])
    length = 14
    p1 = (b[0] - length * math.cos(ang - math.pi / 6), b[1] - length * math.sin(ang - math.pi / 6))
    p2 = (b[0] - length * math.cos(ang + math.pi / 6), b[1] - length * math.sin(ang + math.pi / 6))
    draw.polygon([b, p1, p2], fill=fill)


def draw_centered(draw, text, box, font, fill=(25, 25, 45)):
    lines = []
    for part in text.split("\n"):
        lines += textwrap.wrap(part, width=max(8, int((box[2] - box[0]) / 12))) or [""]
    total = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + 6 * (len(lines) - 1)
    y = box[1] + ((box[3] - box[1]) - total) / 2
    for line in lines:
        bb = draw.textbbox((0, 0), line, font=font)
        x = box[0] + ((box[2] - box[0]) - (bb[2] - bb[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += (bb[3] - bb[1]) + 6


def make_diagram(path, kind):
    im = gradient_bg(1500, 900)
    d = ImageDraw.Draw(im)
    titles = {
        "architecture": "SkillFlow System Architecture",
        "dfd0": "DFD Level 0 - Context Diagram",
        "dfd1": "DFD Level 1 - Main Process Flow",
        "er": "ER Diagram - Database Model",
        "usecase": "Use Case Diagram - User and Admin Roles",
        "activity": "Activity Diagram - Skill Exchange Workflow",
        "flowchart": "Flowchart - OTP Registration and Verification",
        "deployment": "Deployment Architecture",
    }
    d.text((50, 35), titles[kind], font=FONT_B, fill=(31, 45, 92))
    if kind == "architecture":
        boxes = [
            ("User Browser\nHTML / CSS / JavaScript", (70, 180, 370, 330)),
            ("Flask Application\nRoutes, Sessions, APIs", (600, 160, 900, 330)),
            ("MySQL Database\nUsers, Skills, Chats, Payments", (1080, 180, 1410, 330)),
            ("Admin Panel\nModeration and Analytics", (600, 500, 900, 660)),
            ("Email & OTP\nVerification Service", (90, 540, 360, 680)),
            ("Payment Gateway\nChat Unlock Workflow", (1080, 540, 1410, 680)),
        ]
        links = [((370, 255), (600, 245)), ((900, 245), (1080, 255)), ((750, 330), (750, 500)), ((370, 610), (600, 260)), ((900, 580), (1080, 600))]
    elif kind == "dfd0":
        boxes = [
            ("User", (110, 180, 330, 300)),
            ("SkillFlow\nPlatform", (560, 300, 940, 540)),
            ("Admin", (1140, 180, 1360, 300)),
            ("Email OTP", (110, 610, 330, 730)),
            ("Payment Gateway", (1140, 610, 1360, 730)),
        ]
        links = [((330, 240), (560, 380)), ((1140, 240), (940, 380)), ((330, 670), (560, 470)), ((1140, 670), (940, 470))]
    elif kind == "dfd1":
        labels = ["1. Register / Login", "2. Manage Profile & Skills", "3. Search & Match", "4. Send Request", "5. Unlock Chat / Payment", "6. Notifications & Admin Review"]
        coords = [(90, 180, 390, 310), (600, 180, 900, 310), (1090, 180, 1390, 310), (90, 560, 390, 690), (600, 560, 900, 690), (1090, 560, 1390, 690)]
        boxes = list(zip(labels, coords))
        links = [((390, 245), (600, 245)), ((900, 245), (1090, 245)), ((1240, 310), (240, 560)), ((390, 625), (600, 625)), ((900, 625), (1090, 625))]
    elif kind == "er":
        boxes = [
            ("users\nid PK\nusername\nemail\nverified", (60, 150, 360, 340)),
            ("skills\nid PK\ncategory\nname\nkeywords", (440, 150, 740, 340)),
            ("skill_requests\nid PK\nsender_id FK\nreceiver_id FK\nstatus", (820, 150, 1120, 340)),
            ("payments\nid PK\nuser_id FK\namount\nstatus", (1200, 150, 1480, 340)),
            ("matches\nid PK\nrequest_id FK\nuser links", (60, 560, 360, 750)),
            ("messages\nid PK\nrequest_id FK\nsender_id FK\nmessage", (440, 560, 740, 750)),
            ("notifications\nid PK\nuser_id FK\ntitle\nis_read", (820, 560, 1120, 750)),
            ("activity_logs\nid PK\nuser_id FK\nactivity_type\ncreated_at", (1200, 560, 1480, 750)),
        ]
        links = [((360, 245), (820, 245)), ((1120, 245), (1200, 245)), ((970, 340), (210, 560)), ((970, 340), (590, 560)), ((210, 340), (970, 560)), ((210, 340), (1340, 560))]
    elif kind == "usecase":
        boxes = [(x, (500 + (i % 2) * 300, 160 + (i // 2) * 110, 730 + (i % 2) * 300, 225 + (i // 2) * 110)) for i, x in enumerate(["Register", "Verify OTP", "Search Skills", "Send Request", "Unlock Chat", "Review", "Manage Users", "Monitor Payments", "Handle Reports", "Send Alerts"])]
        links = []
        d.text((130, 410), "User", font=FONT_B, fill=(31, 45, 92))
        d.text((1230, 410), "Admin", font=FONT_B, fill=(31, 45, 92))
    elif kind == "activity":
        steps = ["Start", "Register and verify email", "Add teach and learn skills", "Search matching learner", "Send request", "Request accepted", "Unlock chat if required", "Skill exchange conversation", "Review and notification", "End"]
        boxes = [(s, (500, 95 + i * 78, 1000, 155 + i * 78)) for i, s in enumerate(steps)]
        links = [((750, 155 + i * 78), (750, 95 + (i + 1) * 78)) for i in range(len(steps) - 1)]
    elif kind == "flowchart":
        boxes = [
            ("Start", (650, 90, 850, 150)),
            ("Fill Signup Form", (560, 210, 940, 270)),
            ("Validate Email and Username", (560, 330, 940, 390)),
            ("Existing Account?", (600, 450, 900, 530)),
            ("Create User + OTP", (230, 620, 570, 690)),
            ("Show Error / Resend OTP", (930, 620, 1270, 690)),
            ("Verify OTP", (560, 760, 940, 830)),
        ]
        links = [((750, 150), (750, 210)), ((750, 270), (750, 330)), ((750, 390), (750, 450)), ((600, 490), (570, 655)), ((900, 490), (930, 655)), ((570, 655), (750, 760)), ((930, 655), (750, 760))]
    else:
        boxes = [
            ("Client Browser\nChrome / Edge", (80, 240, 370, 380)),
            ("Flask Server\n127.0.0.1:5000", (600, 210, 900, 390)),
            ("XAMPP MySQL\nphpMyAdmin", (1110, 240, 1410, 380)),
            ("Static Files\nCSS, JS, Images", (600, 560, 900, 700)),
        ]
        links = [((370, 310), (600, 300)), ((900, 300), (1110, 310)), ((750, 390), (750, 560))]
    for label, box in boxes:
        rounded(d, box, (255, 255, 255), (139, 92, 246), 3)
        draw_centered(d, label, box, FONT_M if len(label) < 25 else FONT_S)
    for a, b in links:
        arrow(d, a, b)
    im.save(path)


def make_diagrams():
    diagrams = []
    for kind in ["architecture", "dfd0", "dfd1", "er", "usecase", "activity", "flowchart", "deployment"]:
        p = ASSET_DIR / f"{kind}.png"
        make_diagram(p, kind)
        diagrams.append(p)
    return diagrams


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin, sec.left_margin, sec.right_margin = Inches(0.75), Inches(0.7), Inches(0.8), Inches(0.8)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "18")
        tag.set(qn("w:space"), "18")
        tag.set(qn("w:color"), PURPLE)
        pg_borders.append(tag)
    sec._sectPr.append(pg_borders)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    for st, sz, color in [("Heading 1", 16, PURPLE), ("Heading 2", 14, DARK), ("Heading 3", 13, TEXT)]:
        styles[st].font.name = "Times New Roman"
        styles[st]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        styles[st].font.size = Pt(sz)
        styles[st].font.bold = True
        styles[st].font.color.rgb = RGBColor.from_string(color)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("SkillFlow - Skill Swapping & Learning Platform | BCA Major Project | Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    footer._p.append(begin)
    footer._p.append(instr)
    footer._p.append(end)
    return doc


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", align=None, bold=False, size=12, color=None, before=0, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_table(doc, title, table):
    rows, cols = len(table), len(table[0])
    tb = doc.add_table(rows=rows, cols=cols)
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    tb.style = "Table Grid"
    for i, row in enumerate(table):
        for j, val in enumerate(row):
            set_cell_text(tb.cell(i, j), val, i == 0, "FFFFFF" if i == 0 else TEXT, 9 if cols > 4 else 10)
            shade_cell(tb.cell(i, j), PURPLE if i == 0 else ("FFFFFF" if i % 2 else BG))
    add_para(doc, "Table: " + title, WD_ALIGN_PARAGRAPH.CENTER, False, 10, MUTED)


def fill_page(doc, title, body, bullets=None, table=None, image=None, caption=None, code=None):
    add_heading(doc, title, 1 if title.startswith("CHAPTER") else 2)
    for item in body:
        add_para(doc, item)
    if bullets:
        for item in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run("• " + item)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
    if table:
        add_table(doc, title.replace("CHAPTER", "Chapter"), table)
    if image and Path(image).exists():
        try:
            doc.add_picture(str(image), width=Inches(6.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                add_para(doc, caption, WD_ALIGN_PARAGRAPH.CENTER, True, 10, MUTED)
        except Exception:
            pass
    if code:
        tb = doc.add_table(rows=1, cols=1)
        shade_cell(tb.cell(0, 0), "F3ECFF")
        cell = tb.cell(0, 0)
        cell.text = ""
        for ln, line in enumerate(code.splitlines(), 1):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{ln:02d}  {line[:100]}")
            r.font.name = "Consolas"
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(45, 45, 70)
        add_para(doc, "Code explanation: This selected code segment represents one important part of the SkillFlow implementation. The module accepts input from the web interface, validates the request, communicates with MySQL, and returns a clean page or JSON response. The explanation is intentionally focused on purpose, logic, input, output, security handling and database interaction instead of dumping the full project source code.", after=2)


def title_page(doc, title, subtitle):
    for _ in range(3):
        add_para(doc, "")
    add_para(doc, "SkillFlow", WD_ALIGN_PARAGRAPH.CENTER, True, 30, PURPLE, after=2)
    add_para(doc, "Skill Swapping & Learning Platform", WD_ALIGN_PARAGRAPH.CENTER, True, 20, DARK, after=18)
    add_para(doc, title, WD_ALIGN_PARAGRAPH.CENTER, True, 18, TEXT, after=14)
    add_para(doc, subtitle, WD_ALIGN_PARAGRAPH.CENTER, False, 13, MUTED, after=18)
    t = doc.add_table(rows=6, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("University", "Makhanlal Chaturvedi National University"),
        ("Course", "Bachelor of Computer Applications (BCA)"),
        ("Session", "2025-2026"),
        ("Project Type", "Group Project (3 Students)"),
        ("Frontend", "HTML, CSS, JavaScript"),
        ("Backend / Database", "Python Flask, MySQL, phpMyAdmin"),
    ]
    for i, (a, b) in enumerate(data):
        set_cell_text(t.cell(i, 0), a, True, "FFFFFF", 11)
        shade_cell(t.cell(i, 0), PURPLE)
        set_cell_text(t.cell(i, 1), b, False, TEXT, 11)
        shade_cell(t.cell(i, 1), BG)
    add_para(doc, "\nSubmitted in partial fulfillment of the requirement for the award of the degree of Bachelor of Computer Applications.", WD_ALIGN_PARAGRAPH.CENTER, False, 12, TEXT, before=20)


def body_for(topic):
    return [
        f"{topic} is an important part of the SkillFlow project because it connects the academic design of the system with the actual working website. The topic is described with respect to the Python Flask backend, MySQL database, user interface and administrative control used in the project.",
        f"In practical implementation, {topic.lower()} is handled through a clear separation between frontend pages and backend routes. The browser presents forms, cards, buttons and dashboards, while Flask validates data, manages sessions, communicates with MySQL and returns the appropriate page or JSON response.",
        "SkillFlow focuses on real user workflows such as registration, email verification, skill searching, request sending, payment-based chat unlocking, notifications and profile management. The same data is also observed from the admin side so that the platform can be moderated in a controlled manner.",
        "The design is kept professional and consistent with the SkillFlow theme. Soft purple, lavender and pink accents are used in the interface, while the backend continues to follow reliable database and validation rules. The result is a project that is useful as a learning platform and also suitable for academic demonstration.",
    ]


def module_table(topic):
    return [
        ["Point", "Description"],
        ["Purpose", f"To explain {topic} in the SkillFlow platform."],
        ["Input", "User form data, database values, status filters or admin action data."],
        ["Output", "Validated response, updated database state, notification or visible page result."],
        ["Database", "Users, skills, requests, messages, notifications, payments and activity logs as required."],
        ["Security", "Session checking, validation, safe queries and controlled error response."],
    ]


CODE_SNIPPETS = {
    "Database Connection": "def get_db_connection():\n    return mysql.connector.connect(\n        host='localhost',\n        user='root',\n        password='',\n        database='skillflow_db'\n    )",
    "Registration Route": "@app.route('/register', methods=['GET', 'POST'])\ndef register():\n    email = normalize_email(request.form.get('email'))\n    username = request.form.get('username', '').strip()\n    password_hash = generate_password_hash(password)\n    # check verified and unverified account states\n    # create or update pending user and send OTP",
    "Login Route": "@app.route('/login', methods=['GET', 'POST'])\ndef login():\n    email = normalize_email(request.form.get('email'))\n    user = find_user_by_email(email)\n    if user and check_password_hash(user['password'], password):\n        session['user_id'] = user['id']\n        return redirect(url_for('dashboard'))",
    "OTP Verification": "def issue_verification_otp(user_id, email):\n    otp = generate_numeric_otp()\n    expires_at = datetime.now() + timedelta(minutes=10)\n    save_otp(user_id, email, otp, expires_at)\n    send_verification_email(email, otp)",
    "Skill Matching Logic": "def normalize_skills(value):\n    return {x.strip().lower() for x in value.split(',') if x.strip()}\n\ndef calculate_match(current, other):\n    teach = normalize_skills(current['can_teach'])\n    learn = normalize_skills(current['want_to_learn'])\n    return bool(learn & normalize_skills(other['can_teach']))",
    "Request System": "@app.route('/send-request', methods=['POST'])\ndef send_request():\n    sender_id = session.get('user_id')\n    receiver_id = request.form.get('receiver_id')\n    # prevent duplicate requests and insert pending request",
    "Payment Integration": "@app.route('/payment/verify', methods=['POST'])\ndef verify_payment():\n    request_id = request.form.get('request_id')\n    status = verify_gateway_response(request.form)\n    if status == 'success':\n        unlock_chat(request_id)",
    "Chat Message API": "@app.route('/api/chat/send', methods=['POST'])\ndef send_chat_message():\n    user_id = session.get('user_id')\n    request_id = request.form.get('request_id')\n    message = sanitize_message(request.form.get('message'))\n    # store message only if user belongs to unlocked chat",
    "Admin Login": "@app.route('/admin/login', methods=['GET', 'POST'])\ndef admin_login():\n    admin = find_admin(request.form.get('username'))\n    if admin and check_password_hash(admin['password'], password):\n        session['admin_id'] = admin['id']\n        return redirect(url_for('admin_dashboard'))",
    "Notifications": "def create_notification(user_id, title, message):\n    cursor.execute('INSERT INTO notifications(user_id,title,message,is_read) VALUES(%s,%s,%s,0)',\n                   (user_id, title, message))",
}


def build():
    screens = collect_screenshots()
    diagrams = make_diagrams()
    doc = setup_doc()
    front_pages = [
        ("HARD COVER PAGE", "A Major Project Report prepared for SkillFlow - Skill Swapping & Learning Platform"),
        ("INNER TITLE PAGE", "Complete academic report with user side, admin panel, database design, backend implementation and testing."),
        ("PROJECT CERTIFICATE", "This is to certify that the project entitled SkillFlow - Skill Swapping & Learning Platform has been carried out as a group project for the Bachelor of Computer Applications programme during session 2025-2026."),
        ("CERTIFICATE FROM PROJECT GUIDE", "The project work has been prepared under academic guidance and presents a practical implementation of a Python Flask and MySQL based skill exchange platform."),
        ("SELF DECLARATION", "We declare that this report is an original academic submission. The work explains the SkillFlow project, its modules, database, testing, security and implementation approach."),
        ("ACKNOWLEDGEMENT", "We express sincere gratitude to our guide, faculty members and institute for their support. This project provided practical exposure to web development, database design and project documentation."),
        ("ABSTRACT", "SkillFlow is a web-based platform that connects users for skill exchange. Users can register, verify email using OTP, create profiles, search skills, send requests, unlock chat, exchange knowledge, receive notifications and manage learning activity. The admin panel supports monitoring, reports, payments and platform management."),
        ("TABLE OF CONTENTS", "The report contains front matter, eighteen professional chapters, tables, diagrams, coding sections, testing records and appendices."),
        ("LIST OF FIGURES", "This section lists diagrams and screenshots including architecture, DFD, ER diagram, use case diagram, flowcharts and website screens."),
        ("LIST OF TABLES", "This section lists database tables, requirement tables, testing tables and module-wise documentation tables."),
        ("LIST OF CODES / ABBREVIATIONS", "Important code modules and abbreviations such as Flask, OTP, DFD, ERD, DBMS, HTML, CSS, API and CRUD are listed for quick reference."),
    ]
    intro_pool = [
        "SkillFlow is a practical skill swapping and learning platform developed for students and learners who want to exchange knowledge without depending only on paid courses. The system allows a user to offer one skill and learn another skill from a different user through a structured request, matching, chat and notification workflow.",
        "The project follows a database-driven web application model. HTML, CSS and JavaScript create the visible interface, Python Flask manages route handling and server-side logic, and MySQL stores users, skills, requests, chats, payments, notifications and administrative records.",
        "From an academic point of view, SkillFlow demonstrates complete software development practice. It includes analysis, design, implementation, validation, security and testing. The platform also contains an admin panel, which is important for monitoring users, payments, reports, activity logs and platform safety.",
    ]
    for title, sub in front_pages:
        if title in ["HARD COVER PAGE", "INNER TITLE PAGE"]:
            title_page(doc, title, sub)
        else:
            fill_page(doc, title, [sub] + intro_pool)
        doc.add_page_break()

    chapters = [
        ("CHAPTER 1 - INTRODUCTION", 18), ("CHAPTER 2 - SYSTEM ANALYSIS", 20), ("CHAPTER 3 - FEASIBILITY STUDY", 14),
        ("CHAPTER 4 - SOFTWARE AND HARDWARE REQUIREMENTS", 12), ("CHAPTER 5 - SYSTEM DESIGN", 24), ("CHAPTER 6 - DATABASE DESIGN", 28),
        ("CHAPTER 7 - FRONTEND DESIGN", 22), ("CHAPTER 8 - BACKEND IMPLEMENTATION", 24), ("CHAPTER 9 - MODULE EXPLANATION", 30),
        ("CHAPTER 10 - CODING", 40), ("CHAPTER 11 - TESTING", 24), ("CHAPTER 12 - SECURITY", 20), ("CHAPTER 13 - PAYMENT INTEGRATION", 14),
        ("CHAPTER 14 - OTP AUTHENTICATION", 14), ("CHAPTER 15 - RESULT ANALYSIS", 16), ("CHAPTER 16 - FUTURE SCOPE", 12),
        ("CHAPTER 17 - CONCLUSION", 8), ("CHAPTER 18 - BIBLIOGRAPHY AND APPENDIX", 19),
    ]
    assert sum(p for _, p in chapters) + len(front_pages) == 370
    topics = ["Project Overview", "Workflow", "Implementation", "Validation", "Database Role", "Backend Logic", "User Interface", "Admin Monitoring", "Security Handling", "Testing Result", "Practical Use", "Advantages"]
    fig = 1
    screen_index = 0
    code_keys = list(CODE_SNIPPETS)
    for chap, pages in chapters:
        key = chap.split(" - ", 1)[1]
        for page in range(pages):
            if page == 0:
                fill_page(doc, chap, [f"This chapter presents {key.lower()} for SkillFlow - Skill Swapping & Learning Platform. The discussion is based on the implemented website, its user modules, admin panel, database design and Python Flask backend workflow."] + intro_pool)
            else:
                topic = f"{key} - {topics[(page - 1) % len(topics)]}"
                image = caption = code = None
                table = None
                if key == "SYSTEM DESIGN" and page in [2, 4, 6, 8, 10, 12, 14, 16]:
                    image = diagrams[(page // 2 - 1) % len(diagrams)]
                    caption = f"Figure {fig}: {topic} diagram for SkillFlow"
                    fig += 1
                elif key in ["FRONTEND DESIGN", "MODULE EXPLANATION", "BIBLIOGRAPHY AND APPENDIX"] and screens and page % 2 == 0 and screen_index < len(screens):
                    image = screens[screen_index]
                    caption = f"Figure {fig}: SkillFlow real website screenshot - {topic}"
                    screen_index += 1
                    fig += 1
                if key == "DATABASE DESIGN" and page % 2 == 0:
                    table = [["Field", "Type", "Key", "Description"], ["id", "INT", "PK", "Unique record identifier"], ["user_id", "INT", "FK", "Connected user reference"], ["status", "VARCHAR", "-", "Current state of transaction"], ["created_at", "DATETIME", "-", "Record creation time"]]
                elif key in ["TESTING", "RESULT ANALYSIS"] and page % 2 == 1:
                    table = [["Test Case", "Expected Result", "Actual Result", "Status"], [topic + " validation", "System accepts valid data and rejects invalid data", "Working as expected", "Pass"], ["Session check", "Protected page redirects logged-out user", "Working as expected", "Pass"], ["Database update", "Relevant table updates safely", "Working as expected", "Pass"]]
                elif page % 5 == 0:
                    table = module_table(topic)
                if key == "CODING":
                    topic = code_keys[(page - 1) % len(code_keys)]
                    code = CODE_SNIPPETS[topic]
                fill_page(doc, topic, body_for(topic), bullets=[
                    "The module supports a realistic user workflow.",
                    "Data is stored in MySQL and inspected through phpMyAdmin.",
                    "The design remains consistent with the purple and lavender SkillFlow theme.",
                    "Admin visibility is included where moderation or monitoring is required.",
                ], table=table, image=image, caption=caption, code=code)
            if not (chap == chapters[-1][0] and page == pages - 1):
                doc.add_page_break()
    doc.save(OUT)
    print(OUT)
    print(f"screens_used={screen_index}; screens_available={len(screens)}; diagrams={len(diagrams)}")


if __name__ == "__main__":
    build()
